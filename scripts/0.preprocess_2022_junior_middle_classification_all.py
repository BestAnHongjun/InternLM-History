import os
import docx
import json
import xml.etree.ElementTree as ET


def hasImage(par):
    """get all of the images in a paragraph
    :param par: a paragraph object from docx
    :return: a list of r:embed
    """
    ids = []
    root = ET.fromstring(par._p.xml)
    namespace = {
             'a':"http://schemas.openxmlformats.org/drawingml/2006/main", \
             'r':"http://schemas.openxmlformats.org/officeDocument/2006/relationships", \
             'wp':"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}

    inlines = root.findall('.//wp:inline',namespace)
    for inline in inlines:
        imgs = inline.findall('.//a:blip', namespace)
        for img in imgs:
            id = img.attrib['{{{0}}}embed'.format(namespace['r'])]
        ids.append(id)

    if ids:
        return True
    else:
        return False


def one_test(now_subject, now_topic, start_idx, end_idx, doc, all_test):
    test = dict()
    test["subject_id"] = now_subject
    test["topic_id"] = now_topic
    content = ""
    for j in range(start_idx, end_idx):
        if hasImage(doc.paragraphs[j]):
            print("\n【have picture!!!!】\n")
            return
        para_text = doc.paragraphs[j].text
        if "图" in para_text:
            print("\n【have picture!!!!】\n")
            return
        if content == "":
            content = para_text
        else:
            content = content + "##n##" + para_text
    content = content.replace('\t', "##t##")
    content = content.replace(r"\t", "##t##")
    content = content.replace("  ", "")
    content = content.replace("  ", "")
    content = content.replace(" ", " ")
    content = content.replace("\n", "##n##")
    content = content.replace("##n####n##", "##n##")
    content = content.replace("##t####t##", "##t##")
    test["content"] = content
    all_test["test"].append(test)
    print(test)


def main():
    docx_path = os.path.join("datasets", "src", "2022_junior_middle_classification_all.docx")
    doc = docx.Document(docx_path)

    subject_list = []
    topic_list = []
    para_nums = len(doc.paragraphs)

    now_subject = 0
    now_topic = 0
    start_idx = 0
    find_test = False
    ans_flag = False
    ans_style = None

    all_test = dict()
    all_test["test"] = []

    for i in range(para_nums):
        text = doc.paragraphs[i].text
        num = len(doc.paragraphs[i].runs)
        if num < 1:
            continue
        run0_style = str(doc.paragraphs[i].runs[0].font.color.rgb)
        if "专题" in text:
            flag = False
            for zt in range(32):
                find_idx = "专题{:02d}".format(zt + 1)
                if find_idx in text:
                    now_subject_text = text.split(find_idx)[-1].strip()
                    if " " in now_subject_text:
                        continue
                    now_subject = zt + 1
                    if len(subject_list) == zt:
                        subject_list.append(now_subject_text)
                    # print(now_subject, now_subject_text, now_topic, now_topic_text, i)

                    if find_test and ans_flag:
                        end_idx = i
                        one_test(now_subject, now_topic, start_idx, end_idx, doc, all_test)
                        start_idx = i

                    flag = True
                    find_test = False
                    ans_flag = False
                    break
            if flag:
                continue
        if "考点" in text:
            flag = False
            for kd in range(127):
                find_idx = "考点{:02d}".format(kd + 1)
                if find_idx in text:
                    now_topic_text = text.split(find_idx)[-1].strip()
                    if " " in now_topic_text:
                        continue
                    now_topic = kd + 1
                    if len(topic_list) == kd:
                        topic_list.append(now_topic_text)
                    # print(now_subject, now_subject_text, now_topic, now_topic_text, i)

                    if find_test and ans_flag:
                        end_idx = i
                        one_test(now_subject, now_topic, start_idx, end_idx, doc, all_test)
                        start_idx = i

                    flag = True
                    find_test = True
                    ans_flag = False
                    start_idx = i + 1
                    break
            if flag:
                continue

        if not find_test:
            continue

        if "【答案】" in text:
            ans_flag = True
            ans_style = run0_style
            continue

        if not ans_flag:
            continue
        else:
            if ans_style != run0_style:
                end_idx = i
                one_test(now_subject, now_topic, start_idx, end_idx, doc, all_test)
                start_idx = i
                ans_flag = False
                continue

    all_test["subjects"] = subject_list
    all_test["topic"] = topic_list
    all_test_str = json.dumps(all_test, ensure_ascii=False, sort_keys=True, indent=4, separators=(',', ':'))
    all_test_str.replace("},", "},\n")

    os.makedirs(os.path.join("datasets", "middle"), exist_ok=True)
    with open(os.path.join("datasets", "middle", "2022_junior_middle_classification_all.json"), "w", encoding='utf-8') as f:
        f.write(all_test_str)
    print(len(all_test["test"]))


if __name__ == "__main__":
    main()
